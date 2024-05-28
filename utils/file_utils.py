import json
from docx import Document as DocxDocument
import shutil
from fpdf import FPDF
def delete_directory(directory_path):
    """删除目录及其所有内容的函数。"""
    try:
        shutil.rmtree(directory_path)
        print(f"The directory {directory_path} has been deleted.")
    except OSError as e:
        print(f"Error: {e.strerror}")

def json_to_pdf(json_file_path, pdf_file_path):
    # 读取JSON文件
    with open(json_file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # 创建PDF对象
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    # 检查数据是否为字典
    if not isinstance(data, dict):
        raise ValueError("JSON content must be a dictionary")

    # 将JSON数据添加到PDF中
    for key, value in data.items():
        pdf.cell(0, 10, txt=f"{key}: {value}", ln=True)

    # 保存PDF到文件
    pdf.output(pdf_file_path)

def convert_json_to_word(json_file_path:str, word_file_path:str):
    """
    json_file_path(str):json文件路径
    word_file_path(str):保存的word文件路径
    """
    with open(json_file_path, 'r', encoding='utf-8') as json_file:
        data = json.load(json_file)

    doc = DocxDocument()

    try:
        if isinstance(data, dict):
            for key, value in data.items():
                paragraph = doc.add_paragraph()
                paragraph.add_run(f"{key}: ").bold = True
                paragraph.add_run(str(value))
        elif isinstance(data, list):
            for item in data:
                doc.add_paragraph(str(item))
        else:
            doc.add_paragraph(str(data))
    except UnicodeEncodeError as e:
        print(f"Encoding error: {e}")
        # 处理编码错误或替换问题字符的代码

    doc.save(word_file_path)

